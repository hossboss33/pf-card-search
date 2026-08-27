"""Verbatim-true .docx export. Spec §1.5, §8.3 (palette note), §9.4.

Two presets:

* ``house`` (§1.5): Calibri, 1.15 line spacing, no space before/after
  paragraphs. Tag = Heading 4 paragraph, bold 13pt. Cite paragraph 11pt
  with the short cite bold and any [bracketed qualification] runs inside
  the fullcite at 8pt. Body rebuilt from the variant's ``markup_html``:
  <mark> -> the chosen color's WD_COLOR_INDEX highlight, <strong>/<u> ->
  bold/underline, <span class="min"> -> 8pt (6pt when the minimized run
  is a whole paragraph).
* ``verbatim``: standard Verbatim-style sizes — tag Heading 4, cite 11pt
  with the short cite bold, body 11pt, minimized text left plain but
  small (8pt).

Both presets emit pocket/hat/block as Heading 1/2/3 when present on the
exported variant; repeated values collapse outline-style, so two cards
under one block share one Heading 3.

Two rules from §1.5 that override everything else:

* **Cites are never restamped.** The exported cite and fullcite strings
  are byte-identical to the stored ones — original attribution intact, no
  team stamp or suffix ever appended.
* **Highlighting is a setting, not a constant** (§8.3): the caller's
  selected color is applied at export time, from Word's fixed
  WD_COLOR_INDEX base palette — which is what keeps the screen, the
  printout, and the exported file identical.

Round-trip invariant (§9.4, M6): re-parsing an exported document with
carddb.docx_parser yields the same canonical_key, spoken text, and
summary. The body is rebuilt run-for-run from the sanitized markup (the
input only ever contains h1–h4/p/br/u/strong/em/mark/span.min), with
nesting preserved, so the parser's §3.4 run classification reproduces the
stored classes exactly.
"""
from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from .docx_parser import SHORT_CITE_RE
from .normalize import normalize

# §8.3: Word's base highlighter palette. .docx files can only store
# highlights from the fixed WD_COLOR_INDEX palette, so these four names
# are the whole setting.
HIGHLIGHT_COLORS = {
    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "yellow": WD_COLOR_INDEX.YELLOW,
    "blue": WD_COLOR_INDEX.BLUE,
    "turquoise": WD_COLOR_INDEX.TURQUOISE,
}

_FONT = "Calibri"
_BRACKET_RE = re.compile(r"\[[^\]]*\]")  # [qualifications] inside a fullcite


# --- speech math (§9.22) ---------------------------------------------------

def spoken_word_count(spoken: str) -> int:
    """Words in the spoken (highlighted-only) projection."""
    return len((spoken or "").split())


def read_time_str(words: int, wpm: int = 250) -> str:
    """'1:47'-style read time for a word count at a given pace."""
    if wpm <= 0:
        raise ValueError("wpm must be positive, got %r" % (wpm,))
    secs = round(max(0, int(words)) * 60.0 / wpm)
    minutes, seconds = divmod(int(secs), 60)
    return "%d:%02d" % (minutes, seconds)


# --- markup_html -> run segments -------------------------------------------

class _MarkupWalker(HTMLParser):
    """Walk sanitized card markup into per-paragraph run segments.

    The input is sanitize_markup() output, so only h1–h4, p, br, u,
    strong, em, mark, and span (class="min") ever appear. Nesting is
    preserved via depth counters, so <strong><u>text</u></strong> becomes
    a single bold+underline segment, matching the parser's classes.
    Heading text is skipped: tag and pocket/hat/block are written from
    the stored fields, never re-derived from markup.
    """

    def __init__(self):
        HTMLParser.__init__(self, convert_charrefs=True)
        self.paragraphs: List[list] = []  # [[ [cls, text, italic], ...], ...]
        self._cur: Optional[list] = None
        self._skip = 0
        self._bold = self._under = self._mark = self._min = self._ital = 0
        self._spans: List[bool] = []  # True where the span was class="min"

    def handle_starttag(self, tag, attrs):
        if tag in ("h1", "h2", "h3", "h4"):
            self._skip += 1
        elif tag == "p":
            self._flush()
            self._cur = []
        elif tag == "br":
            self._emit("\n")
        elif tag == "strong":
            self._bold += 1
        elif tag == "u":
            self._under += 1
        elif tag == "mark":
            self._mark += 1
        elif tag == "em":
            self._ital += 1
        elif tag == "span":
            cls = next((v for k, v in attrs if k == "class"), None) or ""
            is_min = "min" in cls.split()
            self._spans.append(is_min)
            if is_min:
                self._min += 1

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "h4"):
            self._skip = max(0, self._skip - 1)
        elif tag == "p":
            self._flush()
        elif tag == "strong":
            self._bold = max(0, self._bold - 1)
        elif tag == "u":
            self._under = max(0, self._under - 1)
        elif tag == "mark":
            self._mark = max(0, self._mark - 1)
        elif tag == "em":
            self._ital = max(0, self._ital - 1)
        elif tag == "span":
            if self._spans and self._spans.pop():
                self._min = max(0, self._min - 1)

    def handle_data(self, data):
        self._emit(data)

    def _cls(self) -> str:
        # Same precedence as docx_parser._classify_run (§3.4).
        if self._mark:
            return "mark"
        if self._bold and self._under:
            return "strong_u"
        if self._under:
            return "u"
        if self._bold:
            return "strong"
        if self._min:
            return "min"
        return "plain"

    def _emit(self, text):
        if self._skip or not text:
            return
        if self._cur is None:
            self._cur = []  # tolerate stray text outside <p>
        cls, ital = self._cls(), bool(self._ital)
        if self._cur and self._cur[-1][0] == cls and self._cur[-1][2] == ital:
            self._cur[-1][1] += text
        else:
            self._cur.append([cls, text, ital])

    def _flush(self):
        if self._cur is not None:
            if "".join(s[1] for s in self._cur).strip():
                self.paragraphs.append(self._cur)
            self._cur = None


def _walk_markup(markup_html: str) -> List[list]:
    w = _MarkupWalker()
    w.feed(markup_html or "")
    w.close()
    w._flush()
    return w.paragraphs


def _body_start(paragraphs: List[list], body_text: str) -> int:
    """Index of the first body paragraph within the markup's paragraphs.

    markup_html includes the cite paragraph(s) ahead of the body; the
    exported cite is written from the stored strings instead, so those
    leading paragraphs must be skipped. The body suffix is found by
    normalize()-equality against the stored body_text — the exact
    equivalence class that defines the canonical card, which is what
    makes the §9.4 round trip hold by construction.
    """
    texts = ["".join(seg[1] for seg in p) for p in paragraphs]
    want = normalize(body_text or "")
    for k in range(len(texts) + 1):
        if normalize(" ".join(texts[k:])) == want:
            return k
    # markup and canonical body drifted (e.g. an HF variant with different
    # raw whitespace): fall back to skipping one cite-shaped lead paragraph
    if texts and SHORT_CITE_RE.match(texts[0].strip()):
        return 1
    return 0


def _body_paragraphs(card, variant) -> List[list]:
    markup = variant["markup_html"] if variant is not None else None
    if markup:
        paras = _walk_markup(markup)
        return paras[_body_start(paras, card["body_text"] or ""):]
    # no markup stored (cites_only fidelity, bare canonical row): plain body
    out = []
    for line in (card["body_text"] or "").split("\n"):
        if line.strip():
            out.append([["plain", line, False]])
    return out


# --- document assembly -----------------------------------------------------

def _setup_styles(doc, preset: str) -> None:
    """House (§1.5): Calibri, 1.15 spacing, 0 space before/after.
    Verbatim: same face, single spacing. Tag style is bold 13pt in both."""
    line = 1.15 if preset == "house" else 1.0
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for lvl in (1, 2, 3, 4):
        st = doc.styles["Heading %d" % lvl]
        st.font.name = _FONT
        spf = st.paragraph_format
        spf.line_spacing = line
        spf.space_before = Pt(0)
        spf.space_after = Pt(0)
    h4 = doc.styles["Heading 4"]
    h4.font.bold = True
    h4.font.italic = False
    h4.font.size = Pt(13)


def _add_cite_run(par, text: str, pt: int):
    run = par.add_run(text)
    run.font.size = Pt(pt)
    return run


def _write_cite(doc, cite: Optional[str], fullcite: Optional[str],
                preset: str) -> None:
    """One cite paragraph: bold short cite, then the fullcite VERBATIM.

    The runs concatenate to exactly ``cite + ", " + fullcite`` — the
    stored strings byte-for-byte, nothing appended (§1.5: cites are never
    restamped). House preset additionally drops any [bracketed
    qualification] runs inside the fullcite to 8pt.
    """
    par = doc.add_paragraph()
    if cite:
        run = _add_cite_run(par, cite, 11)
        run.font.bold = True
    if fullcite:
        if cite:
            _add_cite_run(par, ", ", 11)
        if preset == "house":
            pos = 0
            for m in _BRACKET_RE.finditer(fullcite):
                if m.start() > pos:
                    _add_cite_run(par, fullcite[pos:m.start()], 11)
                _add_cite_run(par, m.group(0), 8)
                pos = m.end()
            if pos < len(fullcite):
                _add_cite_run(par, fullcite[pos:], 11)
        else:
            _add_cite_run(par, fullcite, 11)


def _write_body_paragraph(doc, segs: list, preset: str, hl_index) -> None:
    par = doc.add_paragraph()
    texty = [s for s in segs if s[1].strip()]
    whole_min = bool(texty) and all(s[0] == "min" for s in texty)
    min_pt = 6 if (preset == "house" and whole_min) else 8
    for cls, text, ital in segs:
        run = par.add_run(text)  # '\n' from <br> becomes a real w:br
        if ital:
            run.font.italic = True
        if cls == "mark":
            run.font.highlight_color = hl_index
        elif cls == "strong_u":
            run.font.bold = True
            run.font.underline = True
        elif cls == "u":
            run.font.underline = True
        elif cls == "strong":
            run.font.bold = True
        elif cls == "min":
            run.font.size = Pt(min_pt)
        # plain: inherit Normal (Calibri 11pt)


def _default_variant(conn, card_id: int):
    """Deterministic default: the earliest variant that has markup."""
    return conn.execute(
        "SELECT * FROM card_variants WHERE card_id = ? "
        "ORDER BY (markup_html IS NULL OR markup_html = ''), id LIMIT 1",
        (card_id,),
    ).fetchone()


def export_cards(conn, card_ids: List[int], out_path,
                 preset: str = "house", highlight: str = "green",
                 variant_ids: Optional[List[int]] = None) -> Path:
    """Export cards to one .docx at out_path; returns the Path (§9.4).

    Cards are written in the given order. ``variant_ids`` picks which
    disclosure's markup to use for its card (one per card; variants whose
    card is not in card_ids are ignored); other cards use the default
    variant, and a card with no stored markup falls back to its plain
    body text.
    """
    if preset not in ("house", "verbatim"):
        raise ValueError(
            "unknown preset %r (expected 'house' or 'verbatim')" % (preset,))
    if highlight not in HIGHLIGHT_COLORS:
        raise ValueError(
            "unknown highlight %r (expected one of: %s)"
            % (highlight, ", ".join(sorted(HIGHLIGHT_COLORS))))
    if not card_ids:
        raise ValueError("no cards selected for export")
    hl_index = HIGHLIGHT_COLORS[highlight]

    chosen = {}
    for vid in (variant_ids or []):
        row = conn.execute(
            "SELECT * FROM card_variants WHERE id = ?", (vid,)).fetchone()
        if row is None:
            raise ValueError("no such variant id: %s" % (vid,))
        chosen.setdefault(row["card_id"], row)

    doc = Document()
    _setup_styles(doc, preset)

    cur_pocket = cur_hat = cur_block = None
    for cid in card_ids:
        card = conn.execute(
            "SELECT * FROM cards WHERE id = ?", (cid,)).fetchone()
        if card is None:
            raise ValueError("no such card id: %s" % (cid,))
        variant = chosen.get(cid)
        if variant is None:
            variant = _default_variant(conn, cid)

        pocket = variant["pocket"] if variant is not None else None
        hat = variant["hat"] if variant is not None else None
        block = variant["block"] if variant is not None else None
        if pocket and pocket != cur_pocket:
            doc.add_heading(pocket, level=1)
            cur_pocket, cur_hat, cur_block = pocket, None, None
        if hat and hat != cur_hat:
            doc.add_heading(hat, level=2)
            cur_hat, cur_block = hat, None
        if block and block != cur_block:
            doc.add_heading(block, level=3)
            cur_block = block

        doc.add_heading(card["tag"] or "", level=4)
        if card["cite"] or card["fullcite"]:
            _write_cite(doc, card["cite"], card["fullcite"], preset)
        for segs in _body_paragraphs(card, variant):
            _write_body_paragraph(doc, segs, preset, hl_index)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
