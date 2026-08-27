"""Docx parser for open-source speech documents. Spec §1.1–1.3, §3.4.

Two segmentation passes:

1. Style pass — Verbatim heading hierarchy (§1.3): Heading 1/2/3 set the
   current pocket/hat/block; each Heading 4 opens a card that closes at the
   next heading of any level.
2. Direct-formatting fallback pass — triggered when the style pass yields
   0 cards or the heading counts look degenerate (most of the document's
   text sits outside any card). A short paragraph (< ~40 words) whose runs
   are >= 80% bold at >= 12.5pt, followed within 2 paragraphs by a
   short-cite-shaped line, is treated as a tag.

Within a card (§3.4): the first 1–2 paragraphs after the tag are checked
for a short cite; the remainder of those paragraphs is the full cite, out
of which the first http(s) token becomes source_url and the first
date-shaped token near the front becomes source_pub_date (ISO).

Cite detection is wider than the §3.4 short-cite regex alone (which stays
the path of first precedence): a paragraph right after the tag that is
fullcite-shaped — reasonably short (<= ~80 words) and carrying at least
one strong cite signal (an http(s) token, a 19xx/20xx year, or an
access-date marker such as "Accessed"/"DOA"/"retrieved") — also counts
as the cite paragraph. In that case cite stays None, the whole paragraph
becomes the fullcite, and the body starts after it. Wiki-loaded PF cards
mostly ship cite='' with everything in fullcite, so their exports (§9.4)
have fullcite-only cite paragraphs; this rule is what keeps their export
-> re-parse -> re-ingest round trip from minting duplicate canonicals,
and it also matches messy real-world cite formats.

Analytic rule, reconciled: spec §1.3 defines an analytic as a Heading-4
with NO BODY under it, while §3.4's shorthand says "no cite-shaped
paragraph -> analytic". The §1.3 definition governs (the HF dataset
itself maps only null-fulltext rows to analytics), so the two are
reconciled as: no cite paragraph and only a trivial body (< ~40 words)
-> analytic, as before; no cite paragraph but a substantial body
(>= ~40 words) -> evidence card with cite=None/fullcite=None and the
full body.

Run markup precedence (§3.4, exact): highlight -> <mark>; bold+underline ->
<strong><u>; underline -> <u>; bold -> <strong>; font size <= 9pt ->
<span class="min">; else plain. Inherited style formatting (character
style, then paragraph style and its base styles) fills in run properties
python-docx reports as None.

Failures raise ParseFailure; callers record parse_status='failed' and move
on — a bad file must never abort a batch.
"""
from __future__ import annotations

import html as _html
import io
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .ingest import CardRecord
from .sanitize import sanitize_markup


class ParseFailure(Exception):
    """Raised for any unparseable input; message is the reason."""


@dataclass
class ParsedDocument:
    cards: List[CardRecord] = field(default_factory=list)  # ordinal 0..n-1
    warnings: List[str] = field(default_factory=list)
    used_fallback: bool = False


# --- regexes ---------------------------------------------------------------

# Spec §3.4 short-cite regex, transcribed exactly (the starting point).
SHORT_CITE_RE = re.compile(
    r"^\s*(?:[A-Z][\w'’.-]+(?:,? (?:and|&) [A-Z][\w'’.-]+)?"
    r"|[A-Z][\w'’.-]+ et al\.?),? ['’]?\d{2}(?:\d{2})?\b"
)

# Extensions: comma-separated author lists ("Rodgers, Cooper, and Smith 19",
# "Rodgers, Cooper 06"). Two-line cites (short cite + manual line break +
# full cite in one paragraph) need no extra pattern — the anchor still
# matches at the start of the paragraph text.
_MULTI_AUTHOR_CITE_RE = re.compile(
    r"^\s*[A-Z][\w'’.-]+(?:,\s*[A-Z][\w'’.-]+)+"
    r"(?:,?\s*(?:and|&)\s*[A-Z][\w'’.-]+)?,?\s+['’]?\d{2}(?:\d{2})?\b"
)

_CITE_RES = (SHORT_CITE_RE, _MULTI_AUTHOR_CITE_RE)

_URL_RE = re.compile(r"https?://[^\s<>\"'\)\]]+", re.IGNORECASE)

_HEADING_STYLE_RE = re.compile(r"heading\s*(\d+)", re.IGNORECASE)

_RE_MDY = re.compile(r"\b(\d{1,2})[-/](\d{1,2})[-/](\d{4})\b")
_RE_MONTH_D_Y = re.compile(
    r"\b(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|"
    r"dec(?:ember)?)\.?\s+(\d{1,2})(?:st|nd|rd|th)?\s*,?\s+(\d{4})\b",
    re.IGNORECASE,
)
_RE_YEAR = re.compile(r"\b(?:19|20)\d{2}\b")
_MONTH_NUM = {"jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
              "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12}

# Access-date markers ("Accessed 1-6-2020", "DOA 8/1/26", "retrieved ...")
# — strong signals that a paragraph is a cite line, not body text.
_ACCESS_RE = re.compile(
    r"\b(?:accessed|access date|date of access|retrieved)\b", re.IGNORECASE)
_DOA_RE = re.compile(r"\bDOA\b")

_DATE_ZONE = 200          # "near the front" of the fullcite
_FULLCITE_MAX_WORDS = 80  # widened cite detection: cap on a fullcite-only par
_ANALYTIC_MAX_BODY_WORDS = 40  # no-cite blocks with >= this many body words
                               # are evidence cards, not analytics (§1.3)
_FALLBACK_TAG_MAX_WORDS = 40
_FALLBACK_BOLD_FRACTION = 0.8
_FALLBACK_MIN_PT = 12.5
_MIN_RUN_PT = 9.0

_TAG_WRAP = {
    "mark": ("<mark>", "</mark>"),
    "strong_u": ("<strong><u>", "</u></strong>"),
    "u": ("<u>", "</u>"),
    "strong": ("<strong>", "</strong>"),
    "min": ('<span class="min">', "</span>"),
    "plain": ("", ""),
}
_SUMMARY_CLASSES = {"mark", "strong_u", "u", "strong"}  # underlined or stronger

_MAC_SOFFICE_PATHS = ("/Applications/LibreOffice.app/Contents/MacOS/soffice",)


# --- small text helpers ----------------------------------------------------

def _clean(s: str) -> str:
    return " ".join((s or "").split())


def _wc(s: str) -> int:
    return len((s or "").split())


def _cite_match(text: str):
    for rx in _CITE_RES:
        m = rx.match(text or "")
        if m:
            return m
    return None


def extract_source_url(fullcite: Optional[str]) -> Optional[str]:
    """First http(s) token in the full cite (§3.4)."""
    if not fullcite:
        return None
    m = _URL_RE.search(fullcite)
    if not m:
        return None
    return m.group(0).rstrip(".,;:!?")


def extract_pub_date(fullcite: Optional[str]) -> Optional[str]:
    """First date-shaped token near the front of the full cite, as ISO.
    Formats: M-D-YYYY (or M/D/YYYY), 'Month D, YYYY', bare YYYY (§3.4)."""
    if not fullcite:
        return None
    zone = fullcite[:_DATE_ZONE]
    cands = []  # (position, specificity: 0 = full date wins ties, iso)
    for m in _RE_MDY.finditer(zone):
        mo, d, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 1 <= mo <= 12 and 1 <= d <= 31 and 1900 <= y <= 2099:
            cands.append((m.start(), 0, "%04d-%02d-%02d" % (y, mo, d)))
    for m in _RE_MONTH_D_Y.finditer(zone):
        mo = _MONTH_NUM[m.group(1)[:3].lower()]
        d, y = int(m.group(2)), int(m.group(3))
        if 1 <= d <= 31 and 1900 <= y <= 2099:
            cands.append((m.start(), 0, "%04d-%02d-%02d" % (y, mo, d)))
    for m in _RE_YEAR.finditer(zone):
        cands.append((m.start(), 1, m.group(0)))
    if not cands:
        return None
    cands.sort(key=lambda t: (t[0], t[1]))
    return cands[0][2]


def _looks_like_fullcite(text: str) -> bool:
    """Heuristic: does this paragraph read like the full-cite line rather
    than the start of the card body?"""
    if not text:
        return False
    if _URL_RE.search(text):
        return True
    if extract_pub_date(text) is not None:
        return True
    return text.count(",") >= 2 and _wc(text) <= 120


def _is_strong_fullcite(text: str) -> bool:
    """Widened cite detection (see module docstring): does this paragraph
    read like a full cite even with no short-cite prefix? Reasonably short
    AND carrying at least one strong cite signal — an http(s) token, a
    19xx/20xx year, or an access-date marker (Accessed / DOA / retrieved).
    Stricter than _looks_like_fullcite (no comma-count heuristic), because
    this one decides on its own whether a cite exists at all."""
    if not text or _wc(text) > _FULLCITE_MAX_WORDS:
        return False
    return bool(
        _URL_RE.search(text)
        or _RE_YEAR.search(text)
        or _ACCESS_RE.search(text)
        or _DOA_RE.search(text)
    )


# --- effective run formatting (direct + inherited) -------------------------

def _style_chain(style) -> Iterable:
    seen = set()
    st = style
    while st is not None and id(st) not in seen:
        seen.add(id(st))
        yield st
        try:
            st = st.base_style
        except Exception:
            st = None


def _resolve(run, par, attr):
    """run.font.<attr>, else the run's character style chain, else the
    paragraph style chain (inherited style formatting, §3.4 note)."""
    v = getattr(run.font, attr, None)
    if v is not None:
        return v
    chains = []
    try:
        chains.append(run.style)
    except Exception:
        pass
    try:
        chains.append(par.style)
    except Exception:
        pass
    for base in chains:
        for st in _style_chain(base):
            f = getattr(st, "font", None)
            v = getattr(f, attr, None) if f is not None else None
            if v is not None:
                return v
    return None


def _run_bold(run, par) -> bool:
    return bool(_resolve(run, par, "bold"))


def _run_underline(run, par) -> bool:
    # True/False, WD_UNDERLINE members (NONE is falsy), or None.
    return bool(_resolve(run, par, "underline"))


def _run_highlighted(run, par) -> bool:
    # WD_COLOR_INDEX member or None; AUTO (0) is falsy = no highlight.
    return bool(_resolve(run, par, "highlight_color"))


def _run_size_pt(run, par) -> Optional[float]:
    v = _resolve(run, par, "size")
    try:
        return v.pt if v is not None else None
    except Exception:
        return None


def _classify_run(run, par) -> str:
    """§3.4 precedence, exact."""
    if _run_highlighted(run, par):
        return "mark"
    bold = _run_bold(run, par)
    under = _run_underline(run, par)
    if bold and under:
        return "strong_u"
    if under:
        return "u"
    if bold:
        return "strong"
    size = _run_size_pt(run, par)
    if size is not None and size <= _MIN_RUN_PT:
        return "min"
    return "plain"


def _iter_par_runs(par):
    """All runs in document order, including runs inside hyperlinks."""
    try:
        content = list(par.iter_inner_content())
    except Exception:
        return list(par.runs)
    out = []
    for item in content:
        if hasattr(item, "add_break"):          # Run
            out.append(item)
        elif hasattr(item, "runs"):             # Hyperlink
            out.extend(item.runs)
    return out


def _par_plain(par) -> str:
    return "".join(r.text for r in _iter_par_runs(par))


def _render_paragraph(par):
    """-> (plain_text, inner_html, summary_fragments, spoken_fragments).
    Manual line breaks (w:br) surface as '\\n' in run text: kept in the
    plain text, rendered as <br> in the HTML."""
    segs = []  # [class, text], adjacent same-class runs merged
    for run in _iter_par_runs(par):
        text = run.text
        if not text:
            continue
        cls = _classify_run(run, par)
        if segs and segs[-1][0] == cls:
            segs[-1][1] += text
        else:
            segs.append([cls, text])
    plain = "".join(t for _, t in segs)
    html_parts, summary, spoken = [], [], []
    for cls, text in segs:
        open_t, close_t = _TAG_WRAP[cls]
        esc = _html.escape(text).replace("\n", "<br>")
        html_parts.append(open_t + esc + close_t)
        frag = _clean(text)
        if frag:
            if cls in _SUMMARY_CLASSES:
                summary.append(frag)
            if cls == "mark":
                spoken.append(frag)
    return plain, "".join(html_parts), summary, spoken


# --- document stream -------------------------------------------------------

def _heading_level(par) -> Optional[int]:
    try:
        st = par.style
    except Exception:
        return None
    if st is None:
        return None
    for cand in (getattr(st, "name", None), getattr(st, "style_id", None)):
        if cand:
            m = _HEADING_STYLE_RE.match(str(cand).strip())
            if m:
                try:
                    return int(m.group(1))
                except ValueError:
                    return None
    return None


def _iter_table_paragraphs(tbl):
    """Cell paragraphs in order, merged cells visited once. The seen-set
    stores the w:tc elements themselves (not id()s): holding a reference
    keeps each lxml proxy alive, which is what makes identity stable."""
    seen = set()
    for row in tbl.rows:
        for cell in row.cells:
            if cell._tc in seen:
                continue
            seen.add(cell._tc)
            for cp in cell.paragraphs:
                yield cp
            # one level of nested tables: text only, no deeper recursion
            for nested in getattr(cell, "tables", []):
                for nrow in nested.rows:
                    for ncell in nrow.cells:
                        if ncell._tc in seen:
                            continue
                        seen.add(ncell._tc)
                        for cp in ncell.paragraphs:
                            yield cp


def _table_text(tbl) -> str:
    return "\n".join(t for t in (_par_plain(p).strip()
                                 for p in _iter_table_paragraphs(tbl)) if t)


def _build_stream(doc):
    """Flatten the document body, in order, into typed entries:
    ("h", level, text) | ("p", Paragraph, text, wc) | ("t", Table, text, wc).
    """
    stream = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            par = Paragraph(child, doc)
            lvl = _heading_level(par)
            if lvl is not None:
                stream.append(("h", lvl, _clean(_par_plain(par))))
            else:
                text = _par_plain(par)
                stream.append(("p", par, text, _wc(text)))
        elif child.tag == qn("w:tbl"):
            tbl = Table(child, doc)
            text = _table_text(tbl)
            stream.append(("t", tbl, text, _wc(text)))
    return stream


# --- card assembly ---------------------------------------------------------

def _split_cite(entries):
    """Cite detection over the first 1–2 paragraphs after the tag (§3.4).
    entries: rendered dicts. Returns (cite, fullcite, body_start, analytic).

    Precedence: the short-cite regex paths first, then the widened
    fullcite-only detection (_is_strong_fullcite; see module docstring).
    ``analytic=True`` here only means "no cite paragraph found" —
    _build_card downgrades that to an evidence card when a substantial
    body follows (the §1.3/§3.4 reconciliation).
    """
    n = len(entries)
    if n == 0:
        return None, None, 0, True

    def txt(i):
        return entries[i]["text"]

    if not entries[0]["from_table"]:
        m = _cite_match(txt(0))
        if m:
            cite = _clean(m.group(0)).rstrip(",")
            rem = _clean(txt(0)[m.end():].lstrip(" \t\n,;:-–—"))
            body_start = 1
            if (len(rem) < 15 and n > 1 and not entries[1]["from_table"]
                    and not _cite_match(txt(1)) and _looks_like_fullcite(txt(1))):
                rem = _clean((rem + " " + txt(1)))
                body_start = 2
            return cite, (rem or None), body_start, False
    # cite in the second paragraph (a short lead-in line sits between the
    # tag and the cite)
    if n > 1 and not entries[1]["from_table"] and _wc(txt(0)) < 60:
        m = _cite_match(txt(1))
        if m:
            cite = _clean(m.group(0)).rstrip(",")
            rem = _clean(txt(1)[m.end():].lstrip(" \t\n,;:-–—"))
            lead = _clean(txt(0))
            full = _clean(" ".join(x for x in (lead, rem) if x))
            return cite, (full or None), 2, False
    # widened detection: a fullcite-shaped paragraph with no short-cite
    # prefix (wiki-loaded PF cards, messy real-world cites) is still the
    # cite paragraph — cite stays None, the whole line is the fullcite
    if not entries[0]["from_table"] and _is_strong_fullcite(txt(0)):
        return None, (_clean(txt(0)) or None), 1, False
    # same, with a short lead-in line between the tag and the fullcite
    if (n > 1 and not entries[0]["from_table"] and not entries[1]["from_table"]
            and _wc(txt(0)) < 60 and _is_strong_fullcite(txt(1))):
        full = _clean(" ".join(x for x in (_clean(txt(0)), _clean(txt(1))) if x))
        return None, (full or None), 2, False
    return None, None, 0, True


def _build_card(tag, pocket, hat, block, members, warnings) -> Optional[CardRecord]:
    """members: stream entries after the tag, up to the card boundary."""
    entries = []
    has_table = False
    for entry in members:
        if entry[0] == "p":
            plain, phtml, summ, spok = _render_paragraph(entry[1])
            if plain.strip():
                entries.append({"text": plain.strip(), "html": phtml,
                                "summary": summ, "spoken": spok,
                                "from_table": False})
        elif entry[0] == "t":
            has_table = True
            for cp in _iter_table_paragraphs(entry[1]):
                plain, phtml, summ, spok = _render_paragraph(cp)
                if plain.strip():
                    entries.append({"text": plain.strip(), "html": phtml,
                                    "summary": summ, "spoken": spok,
                                    "from_table": True})

    cite, fullcite, body_start, analytic = _split_cite(entries)
    tag = _clean(tag or "")
    if analytic and not tag:
        warnings.append("dropped a card with neither tag nor cite")
        return None

    body_entries = entries[body_start:]
    body_text = "\n".join(e["text"] for e in body_entries)
    if analytic and _wc(body_text) >= _ANALYTIC_MAX_BODY_WORDS:
        # §1.3 governs: a substantial body under the tag means evidence,
        # even when no cite paragraph was recognized (module docstring).
        analytic = False
    if not analytic and not body_text:
        warnings.append(
            "card %r has a cite but no body; treated as analytic" % tag[:40])
        analytic = True
    summary = " ".join(f for e in body_entries for f in e["summary"])
    spoken = " ".join(f for e in body_entries for f in e["spoken"])
    ratio = min(1.0, len(spoken) / len(body_text)) if body_text else 0.0

    html_parts = []
    if tag:
        html_parts.append("<h4>" + _html.escape(tag) + "</h4>")
    for e in entries:
        html_parts.append("<p>" + e["html"] + "</p>")

    rec = CardRecord(
        tag=tag or None,
        cite=cite,
        fullcite=fullcite,
        body_text=body_text,
        is_analytic=analytic,
        source_url=extract_source_url(fullcite),
        source_pub_date=extract_pub_date(fullcite),
        pocket=pocket, hat=hat, block=block,
        markup_html=sanitize_markup("".join(html_parts)),
        summary=summary,
        spoken=spoken,
        highlight_ratio=ratio,
    )
    if has_table:
        rec.extras["has_table"] = True
    return rec


# --- pass 1: heading styles ------------------------------------------------

def _style_pass(stream, warnings):
    pocket = hat = block = None
    raw_cards = []
    current = None
    unassigned_words = 0
    for entry in stream:
        if entry[0] == "h":
            if current is not None:
                raw_cards.append(current)
                current = None
            lvl, text = entry[1], entry[2]
            if lvl == 1:
                pocket, hat, block = (text or None), None, None
            elif lvl == 2:
                hat, block = (text or None), None
            elif lvl == 3:
                block = text or None
            elif lvl == 4:
                current = {"tag": text, "pocket": pocket, "hat": hat,
                           "block": block, "members": []}
            else:
                warnings.append("ignored heading level %d: %r" % (lvl, text[:40]))
        else:
            if current is not None:
                current["members"].append(entry)
            else:
                unassigned_words += entry[3]
    if current is not None:
        raw_cards.append(current)
    return raw_cards, unassigned_words


# --- pass 2: direct-formatting fallback ------------------------------------

def _is_tag_shaped(par, text) -> bool:
    """Short paragraph, >= 80% of its characters in bold runs at >= 12.5pt
    (direct or style-inherited), and not itself cite-shaped."""
    words = _wc(text)
    if words == 0 or words >= _FALLBACK_TAG_MAX_WORDS:
        return False
    if _cite_match(text):
        return False
    total = boldlarge = 0
    for run in _iter_par_runs(par):
        t = run.text
        if not t:
            continue
        total += len(t)
        if _run_bold(run, par):
            size = _run_size_pt(run, par)
            if size is not None and size >= _FALLBACK_MIN_PT:
                boldlarge += len(t)
    return total > 0 and (boldlarge / total) >= _FALLBACK_BOLD_FRACTION


def _fallback_pass(stream, warnings):
    seq = [e for e in stream
           if e[0] == "h" or (e[0] in ("p", "t") and e[2].strip())]
    # which paragraphs are tags: tag-shaped, followed within 2 paragraphs
    # (no heading in between) by a short-cite-shaped line
    tag_idx = set()
    for i, e in enumerate(seq):
        if e[0] != "p" or not _is_tag_shaped(e[1], e[2]):
            continue
        seen_paras = 0
        for j in range(i + 1, len(seq)):
            nxt = seq[j]
            if nxt[0] == "h":
                break
            if nxt[0] != "p":
                continue
            seen_paras += 1
            if _cite_match(nxt[2]):
                tag_idx.add(i)
                break
            if seen_paras >= 2:
                break

    pocket = hat = block = None
    raw_cards = []
    current = None
    for i, e in enumerate(seq):
        if e[0] == "h":
            if current is not None:
                raw_cards.append(current)
                current = None
            lvl, text = e[1], e[2]
            if lvl == 1:
                pocket, hat, block = (text or None), None, None
            elif lvl == 2:
                hat, block = (text or None), None
            elif lvl == 3:
                block = text or None
            elif lvl == 4:
                current = {"tag": text, "pocket": pocket, "hat": hat,
                           "block": block, "members": []}
        elif i in tag_idx:
            if current is not None:
                raw_cards.append(current)
            current = {"tag": _clean(e[2]), "pocket": pocket, "hat": hat,
                       "block": block, "members": []}
        elif current is not None:
            current["members"].append(e)
    if current is not None:
        raw_cards.append(current)
    return raw_cards


# --- top level -------------------------------------------------------------

def _finalize(raw_cards, warnings) -> List[CardRecord]:
    out = []
    for rc in raw_cards:
        rec = _build_card(rc["tag"], rc["pocket"], rc["hat"], rc["block"],
                          rc["members"], warnings)
        if rec is not None:
            out.append(rec)
    return out


def _parse_document(doc) -> ParsedDocument:
    warnings: List[str] = []
    stream = _build_stream(doc)
    total_words = sum(e[3] for e in stream if e[0] in ("p", "t"))

    style_raw, unassigned_words = _style_pass(stream, warnings)
    cards = _finalize(style_raw, warnings)

    used_fallback = False
    run_fallback = (
        (not cards and total_words > 0)
        or (bool(cards) and total_words >= 100
            and unassigned_words / max(total_words, 1) > 0.6)
    )
    if run_fallback:
        fb_cards = _finalize(_fallback_pass(stream, warnings), warnings)
        if len(fb_cards) > len(cards):
            warnings.append(
                "direct-formatting fallback used: %d card(s) from styles, "
                "%d from formatting" % (len(cards), len(fb_cards)))
            cards = fb_cards
            used_fallback = True
        elif not cards:
            used_fallback = True
            warnings.append("no cards found (style and fallback passes)")
        else:
            warnings.append("fallback pass triggered but found no more "
                            "cards; kept style-pass results")

    if not cards and total_words == 0:
        warnings.append("document contains no text")

    for i, rec in enumerate(cards):
        rec.ordinal = i
    return ParsedDocument(cards=cards, warnings=warnings,
                          used_fallback=used_fallback)


def parse_docx_bytes(data: bytes, filename: str = "") -> ParsedDocument:
    """Parse in-memory .docx bytes (wraps BytesIO). Raises ParseFailure."""
    label = filename or "<bytes>"
    if not data:
        raise ParseFailure("empty file: %s has no bytes" % label)
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ParseFailure("not a readable .docx (%s): %s" % (label, exc))
    try:
        return _parse_document(doc)
    except ParseFailure:
        raise
    except Exception as exc:  # never crash on weird content — report it
        raise ParseFailure("parser error in %s: %r" % (label, exc))


def parse_docx(path) -> ParsedDocument:
    """Parse a .docx file from disk. Raises ParseFailure."""
    p = Path(path)
    try:
        data = p.read_bytes()
    except OSError as exc:
        raise ParseFailure("cannot read %s: %s" % (p, exc))
    return parse_docx_bytes(data, filename=p.name)


def convert_doc_to_docx(path) -> Path:
    """Convert a legacy .doc via `soffice --headless --convert-to docx`
    (spec §3.4). Raises ParseFailure when LibreOffice is unavailable or the
    conversion fails."""
    src = Path(path)
    if not src.exists():
        raise ParseFailure("no such file: %s" % src)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        for cand in _MAC_SOFFICE_PATHS:
            if Path(cand).exists():
                soffice = cand
                break
    if soffice is None:
        raise ParseFailure(
            "soffice (LibreOffice) not available; cannot convert %s" % src.name)
    outdir = Path(tempfile.mkdtemp(prefix="carddb-doc2docx-"))
    cmd = [soffice, "--headless", "--convert-to", "docx",
           "--outdir", str(outdir), str(src)]
    try:
        proc = subprocess.run(cmd, capture_output=True, timeout=180)
    except Exception as exc:
        raise ParseFailure("soffice failed for %s: %s" % (src.name, exc))
    out = outdir / (src.stem + ".docx")
    if proc.returncode != 0 or not out.exists():
        err = (proc.stderr or b"").decode("utf-8", "replace").strip()
        raise ParseFailure(
            "soffice conversion failed for %s: rc=%s %s"
            % (src.name, proc.returncode, err[:300]))
    return out
