"""docx_parser tests. Spec §1.1–1.3, §3.4: style pass, fallback pass,
cite/date/url extraction, run-markup precedence, and every §3.4 edge case
with its own programmatic fixture (tests/fixtures/docx_builders.py)."""
import shutil
import sys
import types
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent))

from fixtures.docx_builders import (add_cite, add_run, build_all_highlighted,
                                    build_analytic_with_text, build_empty,
                                    build_loose_pf, build_manual_breaks,
                                    build_non_verbatim, build_table_doc,
                                    build_two_para_cite, build_verbatim,
                                    build_year_only_cite, docx_bytes)

import carddb.docx_parser as dp
from carddb.docx_parser import (ParseFailure, SHORT_CITE_RE,
                                convert_doc_to_docx, extract_pub_date,
                                extract_source_url, parse_docx,
                                parse_docx_bytes)


# --- the spec's short-cite regex (§3.4), transcribed exactly ---------------

@pytest.mark.parametrize("line", [
    "Diamond '13",
    "Kessler '26, Jake Kessler, energy analyst",
    "Rodgers and Cooper 06",
    "Smith et al. 24",
    "Smith et al. '24",
    "Kessler 2026",
    "O'Brien '19",
])
def test_short_cite_re_matches(line):
    assert SHORT_CITE_RE.match(line)


@pytest.mark.parametrize("line", [
    "The grid is failing badly",
    "lowercase name '22",
    "Restricting supply raises prices",
    "",
])
def test_short_cite_re_rejects(line):
    assert SHORT_CITE_RE.match(line) is None


def test_extended_cite_multi_author():
    # extension beyond the spec regex: comma-separated author lists
    assert dp._cite_match("Rodgers, Cooper, and Smith 19, fellows at CFR")
    assert dp._cite_match("Rodgers, Cooper 06")


# --- cite-field extraction helpers -----------------------------------------

@pytest.mark.parametrize("text,want", [
    ("Jake Kessler, 7-14-2026, Wired", "2026-07-14"),
    ("Jake Kessler, 7/14/2026, Wired", "2026-07-14"),
    ("senior fellow, March 12, 2006, Foreign Affairs", "2006-03-12"),
    ("staff writer, Sept. 3, 2024, Vox", "2024-09-03"),
    ("researchers at MIT, 2024, NBER", "2024"),
    ("13-45-2026 is not a date but 2026 is", "2026"),
    ("no dates here at all", None),
    (None, None),
])
def test_extract_pub_date(text, want):
    assert extract_pub_date(text) == want


def test_extract_source_url():
    assert extract_source_url("x, https://example.com/a-b, accessed") == \
        "https://example.com/a-b"
    assert extract_source_url("see http://example.org/p.") == \
        "http://example.org/p"
    assert extract_source_url("no links") is None
    assert extract_source_url(None) is None


# --- style pass on a standard Verbatim file --------------------------------

def test_verbatim_structure_and_context():
    parsed = parse_docx_bytes(docx_bytes(build_verbatim()), "verbatim.docx")
    assert not parsed.used_fallback
    assert len(parsed.cards) == 3
    assert [c.ordinal for c in parsed.cards] == [0, 1, 2]

    c0, c1, c2 = parsed.cards
    assert c0.tag == "Interconnection queues collapsing now"
    assert (c0.pocket, c0.hat, c0.block) == \
        ("Case", "Grid Advantage", "Uniqueness")
    assert c1.tag == "Moratorium kills the grid transition"
    # third card sits under a new Heading-3 block
    assert c2.block == "A2: Data centers good"
    assert (c2.pocket, c2.hat) == ("Case", "Grid Advantage")


def test_verbatim_cite_fields():
    parsed = parse_docx_bytes(docx_bytes(build_verbatim()))
    c0, c1, _ = parsed.cards
    assert c0.cite == "Kessler '26"
    assert c0.fullcite.startswith("Jake Kessler, energy analyst")
    assert c0.source_url == "https://example.com/grid-queue"
    assert c0.source_pub_date == "2026-07-14"  # not the access date
    assert c1.cite == "Diamond '13"
    assert c1.source_pub_date == "2013-03-12"  # Month D, YYYY form
    assert c1.source_url == "https://example.org/collapse-paper"


def test_run_markup_precedence_and_projections():
    parsed = parse_docx_bytes(docx_bytes(build_verbatim()))
    c0 = parsed.cards[0]
    m = c0.markup_html
    assert "<h4>Interconnection queues collapsing now</h4>" in m
    assert '<span class="min">The grid is under strain. </span>' in m
    assert "<u>Interconnection queues have exploded </u>" in m
    assert "<strong><u>beyond any precedent </u></strong>" in m
    # highlight wins over underline: <mark> only, no <u> wrapper
    assert "<mark>and reform is failing</mark>" in m
    assert "<u>and reform is failing" not in m

    assert c0.body_text == ("The grid is under strain. Interconnection "
                            "queues have exploded beyond any precedent "
                            "and reform is failing.")
    # summary = underlined or stronger; minimized text excluded
    assert c0.summary == ("Interconnection queues have exploded beyond "
                          "any precedent and reform is failing")
    assert c0.spoken == "and reform is failing"
    assert c0.highlight_ratio == pytest.approx(
        len(c0.spoken) / len(c0.body_text))
    assert not c0.is_analytic


def test_analytic_heading_with_no_body():
    parsed = parse_docx_bytes(docx_bytes(build_verbatim()))
    c2 = parsed.cards[2]
    assert c2.is_analytic
    assert c2.tag == "Extend: their evidence is outdated"
    assert c2.cite is None and c2.fullcite is None
    assert c2.body_text == ""
    assert c2.highlight_ratio == 0.0


def test_analytic_with_explanation_text():
    parsed = parse_docx_bytes(docx_bytes(build_analytic_with_text()))
    assert len(parsed.cards) == 1
    c = parsed.cards[0]
    assert c.is_analytic
    assert c.cite is None
    assert "interconnection reform" in c.body_text


def test_same_body_different_markup_same_canonical_key():
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    doc = Document()
    doc.add_heading("Interconnection queues collapsing now", level=4)
    add_cite(doc)
    p = doc.add_paragraph()
    add_run(p, "The grid is under strain. ")
    add_run(p, "Interconnection queues have exploded ", bold=True)
    add_run(p, "beyond any precedent ", highlight=WD_COLOR_INDEX.BLUE)
    add_run(p, "and reform is failing", underline=True)
    add_run(p, ".")
    a = parse_docx_bytes(docx_bytes(build_verbatim())).cards[0]
    b = parse_docx_bytes(docx_bytes(doc)).cards[0]
    assert a.body_text == b.body_text
    assert a.key() == b.key()          # §4.2: body governs, markup doesn't
    assert a.markup_html != b.markup_html


# --- edge case: manual line breaks inside a card ---------------------------

def test_manual_line_breaks_inside_card():
    parsed = parse_docx_bytes(docx_bytes(build_manual_breaks()))
    assert len(parsed.cards) == 1
    c = parsed.cards[0]
    # two-line cite paragraph: short cite before the break, fullcite after
    assert c.cite == "Kessler '26"
    assert c.fullcite.startswith("Jake Kessler, energy reporter")
    assert c.source_pub_date == "2026-07-14"
    # the break stays a line break in body_text and becomes <br> in markup
    assert "First sentence of the body" in c.body_text
    assert "second line after a manual break" in c.body_text
    assert "\n" in c.body_text
    assert "<br>" in c.markup_html
    assert c.spoken == "second line after a manual break"
    assert "First sentence of the body" in c.summary


# --- edge case: a table inside a body --------------------------------------

def test_table_inside_body():
    parsed = parse_docx_bytes(docx_bytes(build_table_doc()))
    assert len(parsed.cards) == 1
    c = parsed.cards[0]
    assert c.extras.get("has_table") is True
    for token in ("Region", "PJM", "260 gigawatts pending"):
        assert token in c.body_text
    # paragraphs after the table still belong to the card
    assert "The trend accelerates through 2030." in c.body_text
    assert c.cite == "Kessler '26"


def test_no_table_no_extras_key():
    parsed = parse_docx_bytes(docx_bytes(build_verbatim()))
    assert "has_table" not in parsed.cards[0].extras


# --- edge case: empty file --------------------------------------------------

def test_empty_document_yields_zero_cards():
    parsed = parse_docx_bytes(docx_bytes(build_empty()), "empty.docx")
    assert parsed.cards == []
    assert not parsed.used_fallback
    assert any("no text" in w for w in parsed.warnings)


def test_zero_bytes_raise_parse_failure():
    with pytest.raises(ParseFailure):
        parse_docx_bytes(b"", "zero.docx")


def test_garbage_bytes_raise_parse_failure():
    with pytest.raises(ParseFailure):
        parse_docx_bytes(b"this is definitely not a zip archive", "junk.docx")


def test_missing_path_raises_parse_failure(tmp_path):
    with pytest.raises(ParseFailure):
        parse_docx(tmp_path / "does-not-exist.docx")


# --- edge case: every run highlighted --------------------------------------

def test_every_run_highlighted_ratio_capped():
    parsed = parse_docx_bytes(docx_bytes(build_all_highlighted()))
    assert len(parsed.cards) == 1
    c = parsed.cards[0]
    assert c.highlight_ratio == 1.0          # capped, no crash
    assert c.spoken.startswith("Every single word")
    # highlight still beats bold+underline on the second paragraph
    assert "<mark>and so is this second paragraph</mark>" in c.markup_html
    assert "<strong><u>and so is this second paragraph" not in c.markup_html


# --- edge case: non-Verbatim template (fallback pass) ----------------------

def test_non_verbatim_triggers_fallback():
    parsed = parse_docx_bytes(docx_bytes(build_non_verbatim()))
    assert parsed.used_fallback
    assert len(parsed.cards) == 2
    c0, c1 = parsed.cards
    assert c0.tag == "Fracking bans spike energy prices"
    assert c0.cite == "Loris '19"
    assert c0.source_pub_date == "2019-04-16"
    assert c0.source_url == "https://example.com/fracking"
    # second tag has NO direct formatting: bold 13pt comes from its
    # paragraph style, so inherited style formatting must be consulted
    assert c1.tag == "Price spikes hit low-income families hardest"
    assert c1.cite == "Rodgers and Cooper 06"
    assert c1.source_pub_date == "2006-03-12"
    assert "<mark>Energy burdens are regressive by construction</mark>" \
        in c1.markup_html


def test_fallback_not_used_on_styled_files():
    assert not parse_docx_bytes(docx_bytes(build_verbatim())).used_fallback
    assert not parse_docx_bytes(docx_bytes(build_loose_pf())).used_fallback


# --- edge case: PF-style loose file, Heading 4s only -----------------------

def test_loose_pf_heading4_only():
    parsed = parse_docx_bytes(docx_bytes(build_loose_pf()))
    assert not parsed.used_fallback
    assert len(parsed.cards) == 2
    for c in parsed.cards:
        assert c.pocket is None and c.hat is None and c.block is None
        assert not c.is_analytic
    assert parsed.cards[0].tag == "Fracking bans spike energy prices"
    assert parsed.cards[1].tag == \
        "Price spikes hit low-income families hardest"
    assert parsed.cards[0].spoken == "immediately and durably"


# --- two-paragraph cites and date/url shapes -------------------------------

def test_two_paragraph_cite():
    parsed = parse_docx_bytes(docx_bytes(build_two_para_cite()))
    assert len(parsed.cards) == 1
    c = parsed.cards[0]
    assert c.cite == "Rodgers and Cooper 06"
    assert c.fullcite.startswith("Paul Rodgers and Marcus Cooper")
    assert c.source_pub_date == "2006-03-12"
    assert c.source_url == "https://example.net/hegemony"
    # the fullcite paragraph is not part of the body
    assert c.body_text.startswith("Deterrence claims collapse")


def test_year_only_cite_no_url():
    parsed = parse_docx_bytes(docx_bytes(build_year_only_cite()))
    c = parsed.cards[0]
    assert c.cite == "Smith et al. 24"
    assert c.source_pub_date == "2024"
    assert c.source_url is None


# --- parse_docx (path) vs parse_docx_bytes ---------------------------------

def test_parse_docx_path_matches_bytes(tmp_path):
    data = docx_bytes(build_verbatim())
    p = tmp_path / "case.docx"
    p.write_bytes(data)
    from_path = parse_docx(p)
    from_bytes = parse_docx_bytes(data, "case.docx")
    assert [c.tag for c in from_path.cards] == \
        [c.tag for c in from_bytes.cards]
    assert [c.key() for c in from_path.cards] == \
        [c.key() for c in from_bytes.cards]


# --- .doc conversion path ---------------------------------------------------

def _no_soffice_anywhere():
    return (shutil.which("soffice") is None
            and shutil.which("libreoffice") is None
            and not any(Path(p).exists() for p in dp._MAC_SOFFICE_PATHS))


def test_convert_doc_missing_soffice(monkeypatch, tmp_path):
    monkeypatch.setattr(dp.shutil, "which", lambda name: None)
    monkeypatch.setattr(dp, "_MAC_SOFFICE_PATHS", ())
    src = tmp_path / "legacy.doc"
    src.write_bytes(b"\xd0\xcf\x11\xe0 old word junk")
    with pytest.raises(ParseFailure) as exc:
        convert_doc_to_docx(src)
    assert "soffice" in str(exc.value)


def test_convert_doc_missing_file():
    with pytest.raises(ParseFailure):
        convert_doc_to_docx("/nonexistent/nowhere.doc")


def test_convert_doc_with_fake_soffice(monkeypatch, tmp_path):
    data = docx_bytes(build_loose_pf())
    src = tmp_path / "legacy.doc"
    src.write_bytes(b"\xd0\xcf\x11\xe0 pretend this is a .doc")

    monkeypatch.setattr(dp.shutil, "which",
                        lambda name: "/usr/bin/fake-soffice")

    def fake_run(cmd, capture_output=True, timeout=None):
        outdir = Path(cmd[cmd.index("--outdir") + 1])
        (outdir / (Path(cmd[-1]).stem + ".docx")).write_bytes(data)
        return types.SimpleNamespace(returncode=0, stdout=b"", stderr=b"")

    monkeypatch.setattr(dp.subprocess, "run", fake_run)
    out = convert_doc_to_docx(src)
    assert out.name == "legacy.docx"
    parsed = parse_docx(out)
    assert len(parsed.cards) == 2


def test_convert_doc_soffice_failure(monkeypatch, tmp_path):
    src = tmp_path / "legacy.doc"
    src.write_bytes(b"junk")
    monkeypatch.setattr(dp.shutil, "which",
                        lambda name: "/usr/bin/fake-soffice")

    def fake_run(cmd, capture_output=True, timeout=None):
        return types.SimpleNamespace(returncode=77, stdout=b"",
                                     stderr=b"conversion exploded")

    monkeypatch.setattr(dp.subprocess, "run", fake_run)
    with pytest.raises(ParseFailure) as exc:
        convert_doc_to_docx(src)
    assert "rc=77" in str(exc.value)


@pytest.mark.skipif(_no_soffice_anywhere(),
                    reason="soffice (LibreOffice) not installed")
def test_convert_with_real_soffice(tmp_path):
    # LibreOffice sniffs by content, so real .docx bytes behind a .doc
    # name exercise the genuine shell conversion path end to end.
    src = tmp_path / "legacy.doc"
    src.write_bytes(docx_bytes(build_loose_pf()))
    out = convert_doc_to_docx(src)
    parsed = parse_docx(out)
    assert len(parsed.cards) == 2
