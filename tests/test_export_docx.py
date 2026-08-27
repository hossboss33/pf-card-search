"""export_docx tests. Spec §1.5 (house format; cites never restamped),
§8.3 (highlight color is a user setting from Word's base palette), and
§9.4 / M6 (the round-trip invariant: export -> re-parse -> identical
canonical_key, spoken text, and summary)."""
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

sys.path.insert(0, str(Path(__file__).resolve().parent))

from fixtures.export_builders import (build_multilayer, build_reader,
                                      docx_bytes)
from fixtures.docx_builders import build_verbatim

from carddb.db import open_db
from carddb.docx_parser import parse_docx, parse_docx_bytes
from carddb.export_docx import (HIGHLIGHT_COLORS, export_cards,
                                read_time_str, spoken_word_count)
from carddb.ingest import CardRecord, attach_variant, insert_card
from carddb.keys import sha256_bytes
from carddb.rawstore import record_document


@pytest.fixture()
def conn(tmp_path):
    c = open_db(tmp_path / "cards.sqlite")
    yield c
    c.close()


def ingest_bytes(conn, data, name="fixture.docx"):
    """Run docx bytes through the real ingest path; return
    [(card_id, variant_id, parsed CardRecord), ...] in document order."""
    parsed = parse_docx_bytes(data, name)
    doc_id = record_document(conn, sha256_bytes(data), "test", None, name, None)
    out = []
    for rec in parsed.cards:
        cid, _ = insert_card(conn, rec)
        vid, _ = attach_variant(conn, cid, rec, doc_id, None)
        out.append((cid, vid, rec))
    conn.commit()
    return out


def find_run(doc, text):
    for par in doc.paragraphs:
        for run in par.runs:
            if run.text == text:
                return run
    raise AssertionError("no run with text %r" % text)


def cite_par_after_tag(doc):
    pars = doc.paragraphs
    for i, par in enumerate(pars):
        if par.style.name == "Heading 4":
            return pars[i + 1]
    raise AssertionError("no Heading 4 paragraph in exported doc")


def heading_texts(doc, style_name):
    return [p.text for p in doc.paragraphs if p.style.name == style_name]


# --- palette (§8.3) --------------------------------------------------------

def test_highlight_colors_are_words_base_palette():
    assert HIGHLIGHT_COLORS["green"] == WD_COLOR_INDEX.BRIGHT_GREEN
    assert HIGHLIGHT_COLORS["yellow"] == WD_COLOR_INDEX.YELLOW
    assert HIGHLIGHT_COLORS["blue"] == WD_COLOR_INDEX.BLUE
    assert HIGHLIGHT_COLORS["turquoise"] == WD_COLOR_INDEX.TURQUOISE
    assert set(HIGHLIGHT_COLORS) == {"green", "yellow", "blue", "turquoise"}


# --- speech math (§9.22) ---------------------------------------------------

@pytest.mark.parametrize("spoken,want", [
    ("", 0),
    (None, 0),
    ("one", 1),
    ("a  b\nc\t d", 4),
    ("  padded   words  ", 2),
])
def test_spoken_word_count(spoken, want):
    assert spoken_word_count(spoken) == want


@pytest.mark.parametrize("words,wpm,want", [
    (0, 250, "0:00"),
    (13, 250, "0:03"),
    (250, 250, "1:00"),
    (446, 250, "1:47"),   # the spec's own example pace line (§9.22)
    (625, 250, "2:30"),
    (1500, 250, "6:00"),
    (300, 300, "1:00"),
    (250, 125, "2:00"),   # configurable WPM
    (-5, 250, "0:00"),    # clamps, never negative
])
def test_read_time_str(words, wpm, want):
    assert read_time_str(words, wpm=wpm) == want


def test_read_time_str_default_wpm_is_250():
    assert read_time_str(250) == "1:00"


def test_read_time_str_rejects_nonpositive_wpm():
    with pytest.raises(ValueError):
        read_time_str(100, wpm=0)


# --- THE round-trip test (§9.4 / M6) ---------------------------------------

@pytest.mark.parametrize("preset", ["house", "verbatim"])
@pytest.mark.parametrize("hl", ["green", "yellow", "blue", "turquoise"])
def test_round_trip_export_reparse(conn, tmp_path, preset, hl):
    """Ingest a realistic multi-layer card (plain + underlined +
    bold-underlined + highlighted + minimized runs), export it, re-parse
    the export: identical canonical_key, spoken text, and summary — for
    both presets and every highlight color."""
    rows = ingest_bytes(conn, docx_bytes(build_multilayer()))
    assert len(rows) == 1
    cid, _vid, rec = rows[0]
    # the fixture really exercises all four marked-up layers
    assert rec.spoken and rec.summary and rec.highlight_ratio > 0
    assert '<span class="min">' in rec.markup_html
    assert "<strong><u>" in rec.markup_html

    out = export_cards(conn, [cid], tmp_path / ("rt-%s-%s.docx" % (preset, hl)),
                       preset=preset, highlight=hl)
    assert out == tmp_path / ("rt-%s-%s.docx" % (preset, hl))
    reparsed = parse_docx(out)
    assert len(reparsed.cards) == 1
    assert not reparsed.used_fallback
    got = reparsed.cards[0]

    assert got.key() == rec.key()          # identical canonical key
    assert got.spoken == rec.spoken        # identical spoken text
    assert got.summary == rec.summary      # identical summary
    # cite fidelity survives the trip too (§1.5)
    assert got.tag == rec.tag
    assert got.cite == rec.cite
    assert got.fullcite == rec.fullcite
    assert got.pocket == rec.pocket
    assert got.hat == rec.hat
    assert got.block == rec.block
    assert not got.is_analytic


def test_round_trip_multi_card_document(conn, tmp_path):
    """A whole Verbatim-style doc (two evidence cards + one analytic)
    exports as one file whose re-parse reproduces every card."""
    rows = ingest_bytes(conn, docx_bytes(build_verbatim()))
    assert len(rows) == 3
    out = export_cards(conn, [cid for cid, _, _ in rows],
                       tmp_path / "multi.docx", preset="house")
    reparsed = parse_docx(out)
    assert len(reparsed.cards) == 3
    for (cid, _vid, rec), got in zip(rows, reparsed.cards):
        assert got.key() == rec.key()
        assert got.spoken == rec.spoken
        assert got.summary == rec.summary
        assert got.is_analytic == rec.is_analytic
        assert got.pocket == rec.pocket
        assert got.hat == rec.hat
        assert got.block == rec.block

    # pocket/hat/block become Heading 1/2/3, collapsed outline-style:
    # cards sharing a pocket/hat/block share one heading paragraph.
    d = Document(str(out))
    assert heading_texts(d, "Heading 1") == ["Case"]
    assert heading_texts(d, "Heading 2") == ["Grid Advantage"]
    assert heading_texts(d, "Heading 3") == ["Uniqueness",
                                             "A2: Data centers good"]
    assert len(heading_texts(d, "Heading 4")) == 3


# --- cites are never restamped (§1.5) --------------------------------------

def test_cites_never_restamped(conn, tmp_path):
    """The exported cite/fullcite are byte-identical to the stored
    strings: no team stamp, no suffix, nothing appended."""
    cid, _vid, rec = ingest_bytes(conn, docx_bytes(build_multilayer()))[0]
    out = export_cards(conn, [cid], tmp_path / "stamp.docx", preset="house")
    d = Document(str(out))
    cite_par = cite_par_after_tag(d)
    # exactly the stored strings, joined by ', ' — nothing else
    assert cite_par.text == rec.cite + ", " + rec.fullcite
    assert cite_par.text.endswith(rec.fullcite)
    # no stamp anywhere in the document
    full_text = "\n".join(p.text for p in d.paragraphs)
    assert "Delbarton" not in full_text
    assert "//" not in cite_par.text.replace("https://", "")
    # and the same holds for the verbatim preset
    out2 = export_cards(conn, [cid], tmp_path / "stamp2.docx", preset="verbatim")
    assert cite_par_after_tag(Document(str(out2))).text == \
        rec.cite + ", " + rec.fullcite


# --- house preset formatting (§1.5) ----------------------------------------

def test_house_preset_formatting(conn, tmp_path):
    cid, _vid, rec = ingest_bytes(conn, docx_bytes(build_multilayer()))[0]
    out = export_cards(conn, [cid], tmp_path / "house.docx",
                       preset="house", highlight="green")
    d = Document(str(out))

    # Calibri, 1.15 line spacing, 0 space before/after
    normal = d.styles["Normal"]
    assert normal.font.name == "Calibri"
    assert normal.font.size.pt == 11
    assert normal.paragraph_format.line_spacing == pytest.approx(1.15)
    assert normal.paragraph_format.space_before.pt == 0
    assert normal.paragraph_format.space_after.pt == 0

    # tag: a Heading 4 paragraph, bold 13pt
    h4 = d.styles["Heading 4"]
    assert h4.font.bold is True
    assert h4.font.size.pt == 13
    assert heading_texts(d, "Heading 4") == [rec.tag]

    # cite paragraph: short cite bold 11pt; [bracketed qualification] 8pt
    cite_par = cite_par_after_tag(d)
    short = cite_par.runs[0]
    assert short.text == "Kessler '26"
    assert short.font.bold is True
    assert short.font.size.pt == 11
    bracket = find_run(d, "[energy analyst at the Grid Institute]")
    assert bracket.font.size.pt == 8
    rest = find_run(d, ", 7-14-2026, “The interconnection queue,” Wired, "
                       "https://example.com/grid-queue, accessed 8-1-2026")
    assert rest.font.size.pt == 11

    # body runs: mark -> chosen WD_COLOR_INDEX; strong/u -> bold/underline
    marked = find_run(d, "grid collapse")
    assert marked.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN
    bu = find_run(d, "past every recorded precedent ")
    assert bu.font.bold is True and bu.font.underline is True
    u = find_run(d, "interconnection queues have exploded ")
    assert u.font.underline is True and not u.font.bold
    plain = find_run(d, "Analysts agree that ")
    assert plain.font.size is None
    assert not plain.font.bold and not plain.font.underline
    assert not plain.font.highlight_color

    # minimized: 8pt inline, 6pt when the whole paragraph is minimized
    # (adjacent same-class runs were merged into one segment at parse time)
    assert find_run(d, "Context kept for evidence ethics. ").font.size.pt == 8
    whole = find_run(d, "This whole paragraph is minimized context "
                        "kept only for completeness.")
    assert whole.font.size.pt == 6


def test_verbatim_preset_formatting(conn, tmp_path):
    cid, _vid, rec = ingest_bytes(conn, docx_bytes(build_multilayer()))[0]
    out = export_cards(conn, [cid], tmp_path / "verb.docx",
                       preset="verbatim", highlight="turquoise")
    d = Document(str(out))

    normal = d.styles["Normal"]
    assert normal.font.name == "Calibri"
    assert normal.font.size.pt == 11
    assert normal.paragraph_format.line_spacing == pytest.approx(1.0)

    # tag still Heading 4; cite 11pt with bold short-cite emphasis
    assert d.styles["Heading 4"].font.bold is True
    cite_par = cite_par_after_tag(d)
    assert cite_par.runs[0].font.bold is True
    # no 8pt bracket rule: the fullcite is uniform 11pt
    for run in cite_par.runs[1:]:
        assert run.font.size.pt == 11
    assert cite_par.text == rec.cite + ", " + rec.fullcite

    # body 11pt (inherited), minimized left plain-but-small (8pt), even
    # for a wholly minimized paragraph
    assert find_run(d, "This whole paragraph is minimized context "
                       "kept only for completeness.").font.size.pt == 8
    assert find_run(d, "Context kept for evidence ethics. ").font.size.pt == 8
    marked = find_run(d, "grid collapse")
    assert marked.font.highlight_color == WD_COLOR_INDEX.TURQUOISE
    assert not marked.font.bold


# --- highlight color is a user setting (§8.3) ------------------------------

def test_highlight_color_applied_at_export_time(conn, tmp_path):
    """Same stored card, two exports, two colors: the .docx highlight
    follows the setting while the spoken text is untouched."""
    cid, _vid, rec = ingest_bytes(conn, docx_bytes(build_multilayer()))[0]
    by_color = {}
    for name in ("green", "blue"):
        out = export_cards(conn, [cid], tmp_path / ("hl-%s.docx" % name),
                           preset="house", highlight=name)
        by_color[name] = Document(str(out))
        assert parse_docx(out).cards[0].spoken == rec.spoken
    assert (find_run(by_color["green"], "grid collapse").font.highlight_color
            == WD_COLOR_INDEX.BRIGHT_GREEN)
    assert (find_run(by_color["blue"], "grid collapse").font.highlight_color
            == WD_COLOR_INDEX.BLUE)


# --- variant selection -----------------------------------------------------

def test_variant_ids_pick_which_disclosure_exports(conn, tmp_path):
    """Two teams disclose the same card with different highlighting; the
    export follows the requested variant's markup."""
    rows_a = ingest_bytes(conn, docx_bytes(build_reader(True)), "team-a.docx")
    rows_b = ingest_bytes(conn, docx_bytes(build_reader(False)), "team-b.docx")
    (cid_a, vid_a, rec_a), (cid_b, vid_b, rec_b) = rows_a[0], rows_b[0]
    assert cid_a == cid_b            # same canonical card
    assert vid_a != vid_b            # two variants
    assert rec_a.spoken != rec_b.spoken

    out_default = export_cards(conn, [cid_a], tmp_path / "va.docx")
    assert parse_docx(out_default).cards[0].spoken == rec_a.spoken

    out_b = export_cards(conn, [cid_a], tmp_path / "vb.docx",
                         variant_ids=[vid_b])
    got_b = parse_docx(out_b).cards[0]
    assert got_b.spoken == rec_b.spoken
    assert got_b.key() == rec_a.key()  # body identical either way


def test_card_without_markup_falls_back_to_plain_body(conn, tmp_path):
    """A canonical row with no variant markup (e.g. cites_only fidelity)
    still exports, and the body round-trips to the same key."""
    rec = CardRecord(
        tag="No markup stored for this card",
        cite="Diamond '13",
        fullcite="Jared Diamond, professor of geography at UCLA, Science, 2013",
        body_text="First body paragraph of the card.\nSecond body paragraph.",
        is_analytic=False,
    )
    cid, created = insert_card(conn, rec)
    assert created
    conn.commit()
    out = export_cards(conn, [cid], tmp_path / "plain.docx")
    got = parse_docx(out).cards[0]
    assert got.key() == rec.key()
    assert got.cite == rec.cite
    assert got.fullcite == rec.fullcite
    assert got.spoken == ""
    assert got.summary == ""


def test_no_headings_emitted_when_variant_has_none(conn, tmp_path):
    cid, _vid, _rec = ingest_bytes(conn, docx_bytes(build_reader(True)))[0]
    out = export_cards(conn, [cid], tmp_path / "loose.docx")
    d = Document(str(out))
    assert heading_texts(d, "Heading 1") == []
    assert heading_texts(d, "Heading 2") == []
    assert heading_texts(d, "Heading 3") == []
    assert len(heading_texts(d, "Heading 4")) == 1


# --- input validation ------------------------------------------------------

def test_export_rejects_bad_inputs(conn, tmp_path):
    cid, _vid, _rec = ingest_bytes(conn, docx_bytes(build_multilayer()))[0]
    with pytest.raises(ValueError):
        export_cards(conn, [cid], tmp_path / "x.docx", preset="fancy")
    with pytest.raises(ValueError):
        export_cards(conn, [cid], tmp_path / "x.docx", highlight="purple")
    with pytest.raises(ValueError):
        export_cards(conn, [999999], tmp_path / "x.docx")
    with pytest.raises(ValueError):
        export_cards(conn, [cid], tmp_path / "x.docx", variant_ids=[999999])
    with pytest.raises(ValueError):
        export_cards(conn, [], tmp_path / "x.docx")
