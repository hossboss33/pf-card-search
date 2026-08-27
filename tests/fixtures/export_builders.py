"""Programmatic .docx fixtures for export_docx tests (spec §1.5, §9.4).

Self-contained (no import from docx_builders) so the export tests stay
decoupled from the parser test fixtures. Every document is built in code
with python-docx.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt


def docx_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def add_run(par, text, bold=None, underline=None, size=None, highlight=None):
    r = par.add_run(text)
    if bold is not None:
        r.font.bold = bold
    if underline is not None:
        r.font.underline = underline
    if size is not None:
        r.font.size = Pt(size)
    if highlight is not None:
        r.font.highlight_color = highlight
    return r


# Fullcite with a [bracketed qualification], a pub date, and a URL —
# exercises the house preset's 8pt bracket rule and date/url extraction.
KESSLER_REST = (", Jake Kessler, [energy analyst at the Grid Institute], "
                "7-14-2026, “The interconnection queue,” Wired, "
                "https://example.com/grid-queue, accessed 8-1-2026")


def add_kessler_cite(doc):
    p = doc.add_paragraph()
    add_run(p, "Kessler '26", bold=True, size=12)
    add_run(p, KESSLER_REST, size=10)
    return p


def build_multilayer():
    """One card exercising every §1.2 body layer: plain + underlined +
    bold-underlined + highlighted + minimized runs across three body
    paragraphs — including a manual line break and a wholly-minimized
    paragraph — under a full pocket/hat/block heading stack."""
    doc = Document()
    doc.add_heading("Case", level=1)
    doc.add_heading("Grid Advantage", level=2)
    doc.add_heading("Uniqueness", level=3)

    doc.add_heading("Interconnection queues are collapsing now", level=4)
    add_kessler_cite(doc)

    p = doc.add_paragraph()
    add_run(p, "Context kept for evidence ethics. ", size=8)              # min
    add_run(p, "Analysts agree that ")                                    # plain
    add_run(p, "interconnection queues have exploded ", underline=True)   # u
    add_run(p, "past every recorded precedent ",
            bold=True, underline=True)                                    # strong+u
    add_run(p, "grid collapse", underline=True,
            highlight=WD_COLOR_INDEX.YELLOW)                              # mark
    add_run(p, " is the base case now.", size=8)                          # min

    b = doc.add_paragraph()
    r = add_run(b, "Reform efforts fail ", underline=True)                # u
    r.add_break()                                                         # manual break
    add_run(b, "every docket stalls", underline=True,
            highlight=WD_COLOR_INDEX.YELLOW)                              # mark
    add_run(b, " before the queue clears.")                               # plain

    m = doc.add_paragraph()                                               # whole-min par
    add_run(m, "This whole paragraph is minimized context ", size=8)
    add_run(m, "kept only for completeness.", size=8)
    return doc


def build_reader(first: bool):
    """Two disclosures of the same card: identical body text, different
    highlighting (team A highlights the opening, team B the closing)."""
    doc = Document()
    doc.add_heading("Interconnection queues are collapsing now", level=4)
    add_kessler_cite(doc)
    p = doc.add_paragraph()
    add_run(p, "Queues have exploded ", underline=True,
            highlight=WD_COLOR_INDEX.YELLOW if first else None)
    add_run(p, "past every precedent ", bold=True, underline=True)
    add_run(p, "and grid collapse follows", underline=True,
            highlight=None if first else WD_COLOR_INDEX.YELLOW)
    add_run(p, " within the decade.")
    return doc
