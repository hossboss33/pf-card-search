"""Programmatic .docx fixture builders for docx_parser tests.

No binary fixtures live in the repo: every test document is built in code
with python-docx, mirroring the real-world shapes spec §3.4 names —
Verbatim heading hierarchy, manual line breaks, tables inside bodies,
empty files, fully-highlighted files, non-Verbatim direct formatting, and
loose PF files with Heading 4s only.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt


def docx_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def add_run(par, text, bold=None, underline=None, size=None, highlight=None):
    r = par.add_run(text)
    if bold is not None:
        r.font.bold = bold
    if underline is not None:
        r.font.underline = underline
    if size is not None:
        r.font.size = Pt(size)
    if highlight is not None:
        r.font.highlight_color = highlight
    return r


KESSLER_REST = (", Jake Kessler, energy analyst at the Grid Institute, "
                "7-14-2026, “The interconnection queue,” Wired, "
                "https://example.com/grid-queue, accessed 8-1-2026")

DIAMOND_REST = (", Jared Diamond, professor of geography at UCLA, Science, "
                "March 12, 2013, https://example.org/collapse-paper")


def add_cite(doc, short="Kessler '26", rest=KESSLER_REST):
    p = doc.add_paragraph()
    add_run(p, short, bold=True, size=12)
    if rest:
        add_run(p, rest, size=10)
    return p


def add_mixed_body(doc, highlight=WD_COLOR_INDEX.YELLOW):
    """One body paragraph exercising every §3.4 run class."""
    p = doc.add_paragraph()
    add_run(p, "The grid is under strain. ", size=8)                        # min
    add_run(p, "Interconnection queues have exploded ", underline=True)     # u
    add_run(p, "beyond any precedent ", bold=True, underline=True)          # strong+u
    add_run(p, "and reform is failing", underline=True, highlight=highlight)  # mark wins
    add_run(p, ".", size=8)
    return p


def build_verbatim():
    """Standard Verbatim file: pocket/hat/block headings, two evidence
    cards, one analytic."""
    doc = Document()
    doc.add_heading("Case", level=1)
    doc.add_heading("Grid Advantage", level=2)
    doc.add_heading("Uniqueness", level=3)

    doc.add_heading("Interconnection queues collapsing now", level=4)
    add_cite(doc)
    add_mixed_body(doc)

    doc.add_heading("Moratorium kills the grid transition", level=4)
    add_cite(doc, short="Diamond '13", rest=DIAMOND_REST)
    p = doc.add_paragraph()
    add_run(p, "Transition timelines slip a decade ", underline=True)
    add_run(p, "when construction pauses", bold=True, underline=True)
    add_run(p, ", according to every published model.")

    doc.add_heading("A2: Data centers good", level=3)
    doc.add_heading("Extend: their evidence is outdated", level=4)  # analytic
    return doc


def build_loose_pf():
    """PF-style loose file: Heading 4 tags only, no pockets or hats."""
    doc = Document()
    doc.add_heading("Fracking bans spike energy prices", level=4)
    add_cite(doc, short="Loris '19",
             rest=(", Nicolas Loris, economist at the Heritage Foundation, "
                   "4-16-2019, https://example.com/fracking"))
    p = doc.add_paragraph()
    add_run(p, "Restricting supply raises prices ", underline=True)
    add_run(p, "immediately and durably", bold=True, underline=True,
            highlight=WD_COLOR_INDEX.BRIGHT_GREEN)
    add_run(p, ", across every state that has tried it.")

    doc.add_heading("Price spikes hit low-income families hardest", level=4)
    add_cite(doc, short="Kessler '26")
    p = doc.add_paragraph()
    add_run(p, "Energy burdens are regressive by construction",
            underline=True)
    add_run(p, ", the data show.")
    return doc


def build_manual_breaks():
    """Manual line breaks inside a card: a two-line cite paragraph (short
    cite, break, full cite) and a body paragraph split by breaks."""
    doc = Document()
    doc.add_heading("Queues break under the moratorium", level=4)
    p = doc.add_paragraph()
    r = add_run(p, "Kessler '26", bold=True, size=12)
    r.add_break()
    add_run(p, "Jake Kessler, energy reporter, 7-14-2026, "
               "https://example.com/grid-queue", size=9)
    b = doc.add_paragraph()
    r = add_run(b, "First sentence of the body ", underline=True)
    r.add_break()
    add_run(b, "second line after a manual break", underline=True,
            highlight=WD_COLOR_INDEX.YELLOW)
    r2 = add_run(b, " third fragment kept plain.")
    return doc


def build_table_doc():
    """A table inside a card body: text extracted, has_table flagged."""
    doc = Document()
    doc.add_heading("Data centers strain regional grids", level=4)
    add_cite(doc)
    p = doc.add_paragraph()
    add_run(p, "Regional operators report record queue growth",
            underline=True)
    add_run(p, " across every interconnection.")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Region"
    tbl.cell(0, 1).text = "Queue growth"
    tbl.cell(1, 0).text = "PJM"
    tbl.cell(1, 1).text = "260 gigawatts pending"
    p = doc.add_paragraph()
    add_run(p, "The trend accelerates through 2030.", underline=True)
    return doc


def build_empty():
    """A structurally valid .docx with no content at all."""
    return Document()


def build_all_highlighted():
    """Every body run highlighted (pre-highlighted paste): ratio caps at 1.0."""
    doc = Document()
    doc.add_heading("Everything is read", level=4)
    add_cite(doc, short="Diamond '13", rest=DIAMOND_REST)
    p = doc.add_paragraph()
    add_run(p, "Every single word of this body is highlighted",
            highlight=WD_COLOR_INDEX.YELLOW)
    p = doc.add_paragraph()
    add_run(p, "and so is this second paragraph",
            highlight=WD_COLOR_INDEX.TURQUOISE, underline=True, bold=True)
    return doc


def build_non_verbatim():
    """No heading styles anywhere: tags are direct bold 13pt (card 1) or
    inherit bold 13pt from a custom paragraph style (card 2), so the
    fallback pass must fire and must consult inherited style formatting."""
    doc = Document()
    st = doc.styles.add_style("TagStyle", WD_STYLE_TYPE.PARAGRAPH)
    st.font.bold = True
    st.font.size = Pt(13)

    p = doc.add_paragraph()
    add_run(p, "Fracking bans spike energy prices", bold=True, size=13)
    add_cite(doc, short="Loris '19",
             rest=(", Nicolas Loris, economist at the Heritage Foundation, "
                   "4-16-2019, https://example.com/fracking"))
    p = doc.add_paragraph()
    add_run(p, "Restricting supply raises prices ", underline=True)
    add_run(p, "immediately and durably", bold=True, underline=True)
    add_run(p, ", across every state that has tried it, and utilities "
               "pass through each increment within one billing cycle.")

    p = doc.add_paragraph(style="TagStyle")
    p.add_run("Price spikes hit low-income families hardest")  # no direct fmt
    p = doc.add_paragraph()
    add_run(p, "Rodgers and Cooper 06", bold=True)
    add_run(p, ", Paul Rodgers and Marcus Cooper, senior fellows, "
               "Foreign Affairs, March 12, 2006, "
               "https://example.net/energy-burdens", size=10)
    p = doc.add_paragraph()
    add_run(p, "Energy burdens are regressive by construction",
            underline=True, highlight=WD_COLOR_INDEX.YELLOW)
    add_run(p, ", the data show, and the bottom quintile pays four times "
               "the share of income the top quintile pays.")
    return doc


def build_two_para_cite():
    """Short cite alone in its own paragraph, full cite in the next one."""
    doc = Document()
    doc.add_heading("Hegemony does not deter conflict", level=4)
    p = doc.add_paragraph()
    add_run(p, "Rodgers and Cooper 06", bold=True, size=12)
    p = doc.add_paragraph()
    add_run(p, "Paul Rodgers and Marcus Cooper, senior fellows at the "
               "Council on Foreign Relations, Foreign Affairs, "
               "March 12, 2006, https://example.net/hegemony", size=9)
    p = doc.add_paragraph()
    add_run(p, "Deterrence claims collapse on inspection", underline=True)
    add_run(p, " when the full record is considered.")
    return doc


def build_year_only_cite():
    """Full cite with only a bare year and no URL."""
    doc = Document()
    doc.add_heading("Automation reshapes labor markets", level=4)
    add_cite(doc, short="Smith et al. 24",
             rest=(", researchers at MIT, 2024, “AI and labor,” "
                   "NBER Working Paper"))
    p = doc.add_paragraph()
    add_run(p, "Displacement outpaces reinstatement", underline=True)
    add_run(p, " in every sector the paper measures.")
    return doc


def build_analytic_with_text():
    """Tag-only block: paragraphs follow but none is cite-shaped."""
    doc = Document()
    doc.add_heading("Extend the uniqueness evidence", level=4)
    p = doc.add_paragraph()
    add_run(p, "their card is from before the interconnection reform, "
               "so it cannot account for the new queue rules")
    return doc
